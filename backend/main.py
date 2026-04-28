import os
import shutil
from typing import Optional
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

from services.scoring import generate_final_score
from services.speech_to_text import transcribe_audio
from services.nlp_analysis import analyze_text
from services.speech_features import extract_speech_features
from services.llm_feedback import generate_audio_feedback
from services.answer_accuracy import evaluate_answer_accuracy
from services.resume_screening import extract_resume_text, extract_key_nouns, screen_resume

# Hiring Platform imports
from database import init_db, save_candidate, save_evaluation, get_all_evaluations, get_evaluation, get_rankings, delete_candidate
from services.resume_analyzer import analyze_resume_text
from services.interview_evaluator import evaluate_interview
from services.github_analyzer import analyze_github_repo
from services.video_evaluator import evaluate_video_transcript, evaluate_video_demo
from services.scoring_engine import calculate_combined_score, generate_score_explanations
from services.pdf_report import generate_pdf_report
from services.transcript_refiner import refine_transcript
from services.jd_analyzer import analyze_jd
from services.jd_matcher import compute_jd_match


app = FastAPI(
    title="AI Candidate Evaluator",
    description="Multi-signal AI-powered candidate evaluation platform with JD matching & portfolio analysis",
    version="3.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
async def startup():
    """Initialize database on application startup."""
    await init_db()


def extract_question_from_file(file_path: str) -> str:
    """Extract question text from a PDF or DOCX file."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    elif ext in (".docx", ".doc"):
        from docx import Document
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    else:
        return ""


def extract_text_from_file(file_path: str) -> str:
    """Extract text from PDF, DOCX, or plain text files."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    elif ext in (".docx", ".doc"):
        from docx import Document
        doc = Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    elif ext in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read().strip()
    else:
        # Try reading as plain text
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except Exception:
            return ""


def save_upload(upload: UploadFile) -> str:
    """Save an uploaded file to a temp path and return the path."""
    file_path = f"temp_{upload.filename}"
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(upload.file, buffer)
    return file_path


def cleanup_file(path: str):
    """Safely remove a temp file."""
    if path and os.path.exists(path):
        os.remove(path)


# =============================================
# UNIFIED CANDIDATE EVALUATION ENDPOINT
# =============================================

@app.post("/evaluate-candidate")
async def evaluate_candidate(
    candidate_name: str = Form(...),
    candidate_email: Optional[str] = Form(None),
    github_repo_url: Optional[str] = Form(None),
    question: Optional[str] = Form(None),
    jd_text: Optional[str] = Form(None),
    resume: Optional[UploadFile] = File(None),
    interview_audio: Optional[UploadFile] = File(None),
    video_file: Optional[UploadFile] = File(None),
    question_file: Optional[UploadFile] = File(None),
    jd_file: Optional[UploadFile] = File(None),
):
    """
    Unified candidate evaluation across 4 dimensions + JD matching:
    - Resume (30%): Upload PDF → auto-extract text → AI analysis
    - Interview (30%): Upload audio → Whisper transcription + speech analysis → AI evaluation
    - GitHub (20%): Provide repo URL or profile → GitHub API fetch → AI evaluation (JD-aware)
    - Video (20%): Upload video/audio → Whisper transcription → AI evaluation
    - JD Match: Upload JD → extract requirements → compute role-fit score
    """
    temp_files = []  # Track files to cleanup

    # --- 0. JOB DESCRIPTION: Extract and analyze ---
    jd_content = None
    jd_requirements = None

    # JD from file upload
    if jd_file and jd_file.filename:
        jd_path = save_upload(jd_file)
        temp_files.append(jd_path)
        jd_content = extract_text_from_file(jd_path)

    # JD from text input (fallback)
    if not jd_content and jd_text and jd_text.strip():
        jd_content = jd_text.strip()

    # Analyze JD to get structured requirements
    if jd_content:
        jd_requirements = await analyze_jd(jd_content)

    # --- 1. RESUME: Extract text from uploaded PDF ---
    resume_text = None
    resume_nouns = None
    if resume and resume.filename:
        resume_path = save_upload(resume)
        temp_files.append(resume_path)
        resume_text = extract_resume_text(resume_path)
        resume_nouns = extract_key_nouns(resume_text)

    # --- 2. INTERVIEW AUDIO: Transcribe + speech analysis ---
    interview_transcript = None
    speech_analysis = None
    nlp_analysis = None
    speech_score = None

    if interview_audio and interview_audio.filename:
        interview_audio_path = save_upload(interview_audio)
        temp_files.append(interview_audio_path)

        # Transcribe using Whisper (prime with resume nouns)
        interview_transcript = transcribe_audio(interview_audio_path, resume_nouns)

        # Run speech feature extraction (tone, confidence, flow, clarity)
        speech_analysis = extract_speech_features(interview_audio_path, interview_transcript)

        # Refine transcript using resume context
        interview_transcript = await refine_transcript(interview_transcript, resume_text, candidate_name)

        # Run NLP analysis
        nlp_analysis = analyze_text(interview_transcript)

        # Generate ML-based speech score
        speech_score = generate_final_score(nlp_analysis, speech_analysis)

    # --- 3. VIDEO: Transcribe uploaded video/audio file ---
    video_transcript = None
    video_path = None
    if video_file and video_file.filename:
        video_path = save_upload(video_file)
        temp_files.append(video_path)

        # Transcribe video audio using Whisper
        video_transcript = transcribe_audio(video_path, resume_nouns)

        # Refine video transcript
        video_transcript = await refine_transcript(video_transcript, resume_text, candidate_name)

    # --- 4. QUESTION: Extract from file if provided ---
    question_text = question.strip() if question and question.strip() else None
    if not question_text and question_file and question_file.filename:
        qf_path = save_upload(question_file)
        temp_files.append(qf_path)
        question_text = extract_question_from_file(qf_path)

    # =============================================
    # RUN ALL 4 AI MODULES (GitHub is JD-aware)
    # =============================================

    # Resume Analyzer
    resume_result = await analyze_resume_text(resume_text or "")

    # Interview Evaluator
    interview_result = await evaluate_interview(interview_transcript or "")

    # GitHub Analyzer (JD-aware: passes requirements for repo selection + relevance scoring)
    github_result = await analyze_github_repo(github_repo_url or "", jd_requirements)

    # Video Evaluator (transcript-only, for scoring engine compatibility)
    video_result = await evaluate_video_transcript(video_transcript or "")

    # Full Video Demo Evaluation (transcript + visual frames)
    video_demo_result = None
    if video_path:
        video_demo_result = await evaluate_video_demo(video_path, video_transcript, jd_requirements)

    # =============================================
    # ORIGINAL ANALYSIS (speech coaching + answer accuracy + resume screening)
    # =============================================

    # LLM feedback on speech
    llm_feedback = None
    if speech_analysis and nlp_analysis and speech_score:
        llm_feedback = await generate_audio_feedback(
            nlp_analysis, speech_analysis, speech_score, interview_transcript
        )

    # Answer accuracy
    accuracy_result = None
    if question_text and interview_transcript:
        accuracy_result = await evaluate_answer_accuracy(question_text, interview_transcript)

    # Resume screening
    resume_screening = None
    if resume_text:
        resume_screening = await screen_resume(resume_text)

    # =============================================
    # COMBINE SCORES + JD MATCH + SAVE
    # =============================================

    scores = calculate_combined_score(resume_result, interview_result, github_result, video_result)

    explanations = await generate_score_explanations(
        scores, resume_result, interview_result, github_result, video_result
    )

    # Compute JD match score
    jd_match_result = None
    if jd_requirements:
        jd_match_result = await compute_jd_match(
            jd_requirements, resume_result, interview_result,
            github_result, video_result, scores
        )

    # Build full report
    full_report = {
        "scores": scores,
        "explanations": explanations,
        "module_details": {
            "resume": resume_result,
            "interview": interview_result,
            "github": github_result,
            "video": video_result
        },
        # JD matching
        "jd_requirements": jd_requirements,
        "jd_match": jd_match_result,
        # Original analysis data
        "speech_analysis": speech_analysis,
        "speech_score": speech_score,
        "nlp_analysis": nlp_analysis,
        "llm_feedback": llm_feedback,
        "answer_accuracy": accuracy_result,
        "resume_screening": resume_screening,
        "interview_transcript": interview_transcript,
        "video_transcript": video_transcript,
        "video_demo": video_demo_result,
    }

    # Save to database
    candidate_id = await save_candidate(candidate_name, candidate_email)
    await save_evaluation(candidate_id, scores, full_report)

    # Cleanup temp files
    for f in temp_files:
        cleanup_file(f)

    return {
        "candidate_id": candidate_id,
        "candidate_name": candidate_name,
        "scores": scores,
        "explanations": explanations,
        "module_details": {
            "resume": resume_result,
            "interview": interview_result,
            "github": github_result,
            "video": video_result
        },
        # JD matching
        "jd_requirements": jd_requirements,
        "jd_match": jd_match_result,
        # Original analysis
        "speech_analysis": speech_analysis,
        "speech_score": speech_score,
        "interview_transcript": interview_transcript,
        "video_transcript": video_transcript,
        "llm_feedback": llm_feedback,
        "answer_accuracy": accuracy_result,
        "resume_screening": resume_screening,
        "video_demo": video_demo_result,
    }


# =============================================
# DATA ENDPOINTS
# =============================================

@app.get("/candidates")
async def list_candidates():
    """Get all evaluated candidates with scores."""
    candidates = await get_all_evaluations()
    return {"candidates": candidates}


@app.get("/candidates/{candidate_id}")
async def get_candidate(candidate_id: int):
    """Get detailed evaluation for a specific candidate."""
    evaluation = await get_evaluation(candidate_id)
    if not evaluation:
        return {"error": "Candidate not found"}, 404
    return evaluation


@app.get("/rankings")
async def get_candidate_rankings():
    """Get candidates ranked by overall score."""
    rankings = await get_rankings()
    return {"rankings": rankings}


@app.get("/candidates/{candidate_id}/report")
async def download_report(candidate_id: int):
    """Download PDF evaluation report for a candidate."""
    evaluation = await get_evaluation(candidate_id)
    if not evaluation:
        return {"error": "Candidate not found"}, 404

    pdf_bytes = generate_pdf_report(evaluation)
    candidate_name = evaluation.get("name", "candidate").replace(" ", "_")
    filename = f"evaluation_report_{candidate_name}.pdf"

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


@app.delete("/candidates/{candidate_id}")
async def remove_candidate(candidate_id: int):
    """Delete a candidate and their evaluation data."""
    deleted = await delete_candidate(candidate_id)
    if not deleted:
        return {"error": "Candidate not found"}, 404
    return {"message": "Candidate deleted successfully"}