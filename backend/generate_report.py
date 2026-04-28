import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    run = heading.runs[0]
    run.font.name = 'Times New Roman'
    if level == 1:
        run.font.size = Pt(16)
        run.font.bold = True
    elif level == 2:
        run.font.size = Pt(14)
        run.font.bold = True
    elif level == 3:
        run.font.size = Pt(12)
        run.font.italic = True
    
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

doc = Document()

# Set Margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)

# Front Page
add_paragraph(doc, "AI Candidate Evaluator: Multi-Signal Skill-Based Hiring Platform\n\n", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER).runs[0].font.size = Pt(18)

add_paragraph(doc, "AIML-452 Major Project - Dissertation\n", align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "submitted in partial fulfillment of the requirement\nfor the award of the degree of\n", align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "Bachelor of Technology\nin\nAIML\n\n", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

p = add_paragraph(doc, "Submitted by\n", align=WD_ALIGN_PARAGRAPH.CENTER)
p = add_paragraph(doc, "Shivansh Yadav\n05613211622\n\n", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

add_paragraph(doc, "Under the supervision of\n", align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "Ms. Kritika Kasera\nAssistant Professor\n\n", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

add_paragraph(doc, "Department of AIML\nGuru Tegh Bahadur Institute of Technology\nNew Delhi\nMay/June 2026", align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# Declaration
add_heading(doc, "DECLARATION", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
decl_text = "This is to certify that the material embodied in this Major Project - Dissertation titled “AI Candidate Evaluator: Multi-Signal Skill-Based Hiring Platform” being submitted in the partial fulfillment of the requirements for the award of the degree of Bachelor of Technology AIML is based on my original work. It is further certified that this Major Project - Dissertation work has not been submitted in full or in part to this university or any other university for the award of any other degree or diploma. My indebtedness to other works has been duly acknowledged at the relevant places.\n\n\nShivansh Yadav\n05613211622"
add_paragraph(doc, decl_text)
doc.add_page_break()

# Certificate
add_heading(doc, "CERTIFICATE", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
cert_text = "This is to certify that the work embodied in this Major Project - Dissertation titled “AI Candidate Evaluator: Multi-Signal Skill-Based Hiring Platform” being submitted in the partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in AIML, is original and has been carried out by Shivansh Yadav (Enrollment No. 05613211622) under my supervision and guidance.\n\nIt is further certified that this Major Project - Dissertation work has not been submitted in full or in part to this university or any other university for the award of any other degree or diploma to the best of my knowledge and belief.\n\n\nMs. Kritika Kasera\nAssistant Professor\n\n\nDr. Pankaj Sharma\nHOD AIML\nGuru Tegh Bahadur Institute of Technology"
add_paragraph(doc, cert_text)
doc.add_page_break()

# Acknowledgement
add_heading(doc, "ACKNOWLEDGEMENT", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
ack_text = "I would like to express my profound gratitude to my faculty supervisor, Ms. Kritika Kasera, for her continuous guidance, encouragement, and invaluable feedback throughout the development of this major project. \n\nI am also deeply thankful to Dr. Pankaj Sharma, HOD of AIML, and the faculty of Guru Tegh Bahadur Institute of Technology, for providing the necessary resources and an environment conducive to learning and innovation.\n\nFinally, I would like to thank my peers, family, and friends for their unwavering support."
add_paragraph(doc, ack_text)
doc.add_page_break()

# Abstract
add_heading(doc, "ABSTRACT", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
abs_text = "The recruitment and hiring landscape is undergoing a massive transformation driven by Artificial Intelligence (AI) and Machine Learning (ML). Traditional hiring processes rely heavily on manual resume screening and subjective interviews, which are often time-consuming, prone to human bias, and fail to capture a holistic view of a candidate's true capabilities.\n\nThis project, the \"AI Candidate Evaluator,\" introduces a comprehensive, multi-signal candidate evaluation system designed to automate and standardize the technical hiring process. By leveraging advanced Natural Language Processing (NLP), Large Language Models (LLMs), Computer Vision, and Audio Processing techniques, the platform evaluates candidates across four distinct dimensions simultaneously: Resume, Interview Transcript and Audio, GitHub Portfolio, and Video Project Demonstrations.\n\nThe system utilizes state-of-the-art technologies such as Whisper for audio transcription, Groq Llama 3 for intelligent answer evaluation, and Vision-capable models to analyze visual UI components. By generating a weighted, composite score across these dimensions, the platform provides recruiters with an objective, data-driven, and highly detailed assessment report."
add_paragraph(doc, abs_text)
doc.add_page_break()

# Content Sections
chapters = [
    ("CHAPTER 1: INTRODUCTION", "The recruitment and hiring landscape is undergoing a massive transformation driven by Artificial Intelligence (AI) and Machine Learning (ML). Traditional hiring processes rely heavily on manual resume screening and subjective interviews, which are often time-consuming, prone to human bias, and fail to capture a holistic view of a candidate's true capabilities.\n\nThis project, the AI Skill-Based Hiring Platform, introduces a comprehensive, multi-signal candidate evaluation system designed to automate and standardize the technical hiring process. By leveraging advanced Natural Language Processing (NLP), Large Language Models (LLMs), Computer Vision, and Audio Processing techniques, the platform evaluates candidates across four distinct dimensions simultaneously: Resume, Interview Transcript and Audio, GitHub Portfolio, and Video Project Demonstrations.\n\nThe system utilizes state-of-the-art technologies such as Whisper for audio transcription, Llama 3 for intelligent answer evaluation and job description matching, and Vision-capable models to analyze visual UI components and code visibility in project demos. By generating a weighted, composite score across these dimensions, the platform provides recruiters with an objective, data-driven, and highly detailed assessment report."),
    
    ("CHAPTER 2: PROBLEM STATEMENT", ""),
    ("2.1 Problem Definition", "The current technical recruitment process suffers from several critical inefficiencies and biases:\n1. Resume Dependency and Keyword Bias: Initial screening is heavily dependent on resumes, leading to unqualified candidates passing while skilled candidates are rejected.\n2. Subjective Interview Evaluations: Human interviewers are subject to conscious and unconscious biases.\n3. Fragmented Assessment Data: Recruiters manually review PDFs, audios, and videos without a centralized scoring system.\n4. Time-Consuming Technical Reviews: Manually reviewing GitHub repos and video demos takes massive engineering time.", 2),
    ("2.2 Objectives", "The primary objective of this project is to build an end-to-end, AI-powered evaluation system. Specific objectives are:\n1. Develop a Multi-Signal Evaluation Engine.\n2. Implement Automated Audio and Speech Analysis.\n3. Integrate Vision and NLP for Portfolio Assessment.\n4. Create an Objective Weighted Scoring System.\n5. Enable Job Description (JD) Matching.", 2),
    
    ("CHAPTER 3: ANALYSIS", ""),
    ("3.1 Software Requirement Specifications", "Functional Requirements:\n1. Multi-Modal Data Ingestion: Allow PDF, MP3, MP4, and GitHub URLs.\n2. Automated Transcription: Transcribe speech via STT models.\n3. AI-Powered Resume and JD Matching: Extract skills and match against JD.\n4. Speech and Tone Analysis: Analyze pitch, WPM, and vocal projection.\n5. Codebase Evaluation: Fetch GitHub data and use LLMs for code quality.\n6. Video Demo Visual Analysis: Extract video frames for Vision models.\n7. Unified Scoring and Reporting: Aggregate scores into weighted final ranking.\n\nNon-Functional Requirements:\n1. Performance and Speed: Asynchronous processing via FastAPI.\n2. Security and Data Privacy: Temporary local file processing only.\n3. Usability and Aesthetics: Modern React glassmorphism UI.\n4. Scalability: Modular AI components.", 2),
    ("3.2 Feasibility Study", "1. Technical Feasibility: Utilizes fast Groq API inferences and robust open-source libraries like Whisper and Librosa.\n2. Economic Feasibility: Low cost due to open-source frameworks and API developer tiers.\n3. Operational Feasibility: Automates manual recruiting operations efficiently.", 2),
    ("3.3 Tools / Technologies / Platform", "Frontend: React.js (JavaScript)\nBackend: FastAPI (Python)\nDatabase: SQLite (aiosqlite)\nLLMs: Groq API (Llama-3.3-70b, Llama-3.2-90b-vision)\nAudio Processing: OpenAI Whisper, Librosa\nVideo Processing: FFmpeg\nDocument Processing: PyPDF2, python-docx, ReportLab", 2),
    
    ("CHAPTER 4: DESIGN AND ARCHITECTURE", ""),
    ("4.1 Structure Chart", "The system architecture is broken down into four tiers:\n1. Presentation Tier (Frontend): React UI\n2. API & Routing Tier (Backend): FastAPI endpoints\n3. AI Service Tier (Core Logic): Isolated Python modules for AI\n4. Data Tier (Database): SQLite managed by aiosqlite", 2),
    ("4.2 Explanation of Modules", "1. Resume Analyzer Module (30% Weight): Extracts text and uses LLMs to score against JD.\n2. Interview Speech Evaluator (30% Weight): Whisper STT and Librosa for communication scores.\n3. GitHub Portfolio Analyzer (20% Weight): Fetches repo stats and uses LLMs for quality scoring.\n4. Video Demo Vision Analyzer (20% Weight): Extracts frames using FFmpeg and scores UI via Vision LLM.\n5. Weighted Scoring Engine: Aggregates the signals into one composite score.", 2),
    ("4.3 Flow Chart", "1. Start: Recruiter uploads Candidate Files and JD.\n2. Parallel Processing: Resume parsing, Audio transcription, Video frame extraction, GitHub fetch.\n3. AI Inference: LLM and Vision models analyze data.\n4. Aggregation: Weighted scoring calculated.\n5. Persistence: Saved to SQLite.\n6. End: Results displayed on React Dashboard.", 2),
    ("4.4 ER Diagram", "Primary tables include 'candidates' (id, name, email) and 'evaluations' (id, candidate_id, scores, full_report) linked by a One-to-One/Many relationship with cascade deletion.", 2),
    
    ("CHAPTER 5: IMPLEMENTATION", "5.1 Screenshots\nImplementation features modern interfaces such as the Unified Upload Interface, Live Processing Loading states with glassmorphism, and the Comprehensive Candidate Scorecard detailing individual AI explanations.\n\n5.2 Source Code\nCore implementations utilize async functions bridging audio and vision models to process video demos concurrently without blocking the server."),
    
    ("CHAPTER 6: TESTING", "Testing ensures robustness:\nTest Case 1: Complete Multi-Signal Upload successfully processes in < 45s.\nTest Case 2: Missing Data dynamically handles missing GitHub/Video inputs by re-weighting.\nTest Case 3: Invalid GitHub Repositories safely return a 0 score catching 404 errors.\nTest Case 4: Corrupt Audio Files gracefully error out without crashing the application pipeline."),
    
    ("CHAPTER 7: SUMMARY AND CONCLUSION", "The AI Candidate Evaluator proves that the recruitment process can be automated, objective, and deeply comprehensive. Integrating Llama 3 models allows the system to perform complex tasks previously requiring human intelligence, ensuring a merit-based, skill-first hiring process."),
    
    ("CHAPTER 8: LIMITATIONS AND FUTURE WORK", "") ,
    ("8.1 Limitations", "Dependencies on external APIs (Groq) create rate limiting risks. Extracting frames and handling high-resolution video imposes significant computational and memory overhead locally.", 2),
    ("8.2 Future Work", "Future iterations aim to include Live Interview Agent Integration acting as a real-time conversational agent, and Interactive Code Execution Environments to run code logic alongside static repository analysis.", 2)
]

for item in chapters:
    title = item[0]
    content = item[1]
    level = item[2] if len(item) > 2 else 1
    
    if title.startswith("CHAPTER"):
        doc.add_page_break()
    
    add_heading(doc, title, level=level)
    if content:
        add_paragraph(doc, content)

doc.save('../Shivansh_Yadav_Dissertation.docx')
print("Document generated successfully.")
